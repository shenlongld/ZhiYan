"""
附件：从网页中识别 PDF/Docx/Xlsx 链接，异步下载并解析为 Markdown，供 LLM 使用。
"""
from __future__ import annotations

import logging
import re
import uuid
from pathlib import Path
from typing import List, Tuple
from urllib.parse import urljoin, urlparse

import html2text
import httpx
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

# 项目根目录下的 temp_files（与 postgrad_agent 同级）
TEMP_DIR = Path(__file__).resolve().parent.parent / "temp_files"

CONTEXT_PREAMBLE = (
    "以下内容包含网页正文及相关附件（PDF/Docx/Xlsx）的解析结果，"
    "请综合分析名额、政策和门槛。\n\n"
)


def ensure_temp_dir() -> Path:
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    return TEMP_DIR


def scan_attachment_urls(html_or_text: str, base_url: str) -> List[str]:
    """从 HTML 或纯文本中提取 .pdf / .docx / .xlsx 的绝对 URL。"""
    if not html_or_text or not base_url:
        return []
    urls: List[str] = []
    if "<" in html_or_text and ">" in html_or_text:
        try:
            soup = BeautifulSoup(html_or_text, "html.parser")
            for a in soup.find_all("a", href=True):
                href = (a.get("href") or "").strip()
                if re.search(r"\.(pdf|docx|xlsx)(\?|#|$)", href, re.I):
                    urls.append(urljoin(base_url, href))
        except Exception as e:
            logger.warning("BeautifulSoup 解析失败，回退正则: %s", e)
    for m in re.finditer(
        r"https?://[^\s<>\"']+\.(?:pdf|docx|xlsx)(?:\?[^\s<>\"']*)?",
        html_or_text,
        re.I,
    ):
        urls.append(m.group(0))
    for m in re.finditer(
        r'href\s*=\s*["\']([^"\']+\.(?:pdf|docx|xlsx)(?:\?[^"\']*)?)["\']',
        html_or_text,
        re.I,
    ):
        urls.append(urljoin(base_url, m.group(1).strip()))
    seen: set[str] = set()
    out: List[str] = []
    for u in urls:
        if u not in seen:
            seen.add(u)
            out.append(u)
    return out


async def download_attachment(
    client: httpx.AsyncClient, url: str, dest_dir: Path
) -> Path | None:
    try:
        r = await client.get(
            url,
            follow_redirects=True,
            timeout=120.0,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            },
        )
        r.raise_for_status()
        name = urlparse(url).path.split("/")[-1] or "download.bin"
        if "?" in name:
            name = name.split("?")[0]
        name = re.sub(r"[^\w.\-]", "_", name)[:120]
        path = dest_dir / f"{uuid.uuid4().hex[:8]}_{name}"
        path.write_bytes(r.content)
        logger.info("已下载附件: %s -> %s", url, path.name)
        return path
    except Exception as e:
        logger.warning("附件下载失败 %s: %s", url, e)
        return None


def html_to_markdown(html: str) -> str:
    if not html or not html.strip():
        return ""
    h = html2text.HTML2Text()
    h.ignore_links = False
    h.body_width = 0
    try:
        return h.handle(html)
    except Exception as e:
        logger.warning("html2text 失败: %s", e)
        return html


def _table_to_md(rows: List[List[str]]) -> str:
    if not rows:
        return ""
    header = [str(c or "").strip().replace("\n", " ") for c in rows[0]]
    if not any(header):
        return ""
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * len(header)) + " |",
    ]
    for row in rows[1:]:
        cells = [str(c or "").strip().replace("\n", " ") for c in row]
        while len(cells) < len(header):
            cells.append("")
        lines.append("| " + " | ".join(cells[: len(header)]) + " |")
    return "\n".join(lines)


def pdf_to_markdown(path: Path) -> str:
    import fitz  # PyMuPDF
    import pdfplumber

    text_parts: List[str] = []
    doc = fitz.open(path)
    try:
        for page in doc:
            text_parts.append(page.get_text())
    finally:
        doc.close()
    plain = "\n".join(text_parts)
    table_md: List[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables() or []
                for table in tables:
                    if table and table[0]:
                        str_rows = [
                            [
                                str(c).strip().replace("\n", " ")
                                if c is not None
                                else ""
                                for c in row
                            ]
                            for row in table
                        ]
                        table_md.append(_table_to_md(str_rows))
    except Exception as e:
        logger.warning("pdfplumber 表格提取失败: %s", e)
    if table_md:
        return plain + "\n\n## PDF 表格（Markdown）\n\n" + "\n\n".join(table_md)
    return plain


def docx_to_markdown(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
        if rows:
            parts.append(_table_to_md(rows))
    return "\n\n".join(parts)


def xlsx_to_markdown(path: Path) -> str:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts: List[str] = []
        for sheet in wb.worksheets:
            data = list(sheet.iter_rows(values_only=True))
            if not data:
                continue
            str_rows = [
                [str(c) if c is not None else "" for c in row] for row in data
            ]
            parts.append(_table_to_md(str_rows))
        wb.close()
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("xlsx 解析失败: %s", e)
        return ""


def parse_attachment_file(path: Path) -> str:
    suf = path.suffix.lower()
    try:
        if suf == ".pdf":
            return pdf_to_markdown(path)
        if suf == ".docx":
            return docx_to_markdown(path)
        if suf == ".xlsx":
            return xlsx_to_markdown(path)
    except Exception as e:
        logger.warning("附件解析失败 %s: %s", path, e)
    return ""


def merge_web_and_attachments(
    web_md: str, attachment_sections: List[Tuple[str, str]]
) -> str:
    parts: List[str] = [CONTEXT_PREAMBLE, "## 网页正文\n", web_md or ""]
    for name, md in attachment_sections:
        if md.strip():
            parts.append(f"\n\n## 附件: {name}\n{md}")
    return "".join(parts)


def truncate_for_llm(text: str, max_tokens: int = 10000) -> str:
    """中文偏多时约 2 字符/token；优先保留含表格片段。"""
    max_chars = max_tokens * 2
    if len(text) <= max_chars:
        return text
    table_blocks: List[str] = []
    for m in re.finditer(r"(?:^\|.+\|\s*\n)+", text, re.MULTILINE):
        table_blocks.append(m.group(0))
    tables_text = "\n\n".join(table_blocks)
    head_budget = max(4000, max_chars // 2 - len(tables_text))
    head = text[:head_budget]
    combined = head + "\n\n## [保留] Markdown 表格片段\n\n" + tables_text
    if len(combined) > max_chars:
        combined = combined[:max_chars] + "\n\n...[已截断]"
    return combined


def safe_unlink_paths(paths: List[Path]) -> None:
    for p in paths:
        try:
            if p.is_file():
                p.unlink()
                logger.debug("已删除临时文件: %s", p)
        except OSError as e:
            logger.warning("删除临时文件失败 %s: %s", p, e)
